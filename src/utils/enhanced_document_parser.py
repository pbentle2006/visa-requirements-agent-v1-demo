import re
import io
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel parsing
try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class EnhancedDocumentParser:
    """Enhanced document parser supporting multiple file formats."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['.txt', '.md']
        
        if PDF_AVAILABLE:
            self.supported_formats.extend(['.pdf'])
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx', '.doc'])
        if EXCEL_AVAILABLE:
            self.supported_formats.extend(['.xlsx', '.xls'])
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return self.supported_formats
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from file with format detection.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with document content and metadata
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_extension = path.suffix.lower()
        
        if file_extension in ['.txt', '.md']:
            return self._load_text_document(path)
        elif file_extension == '.pdf' and PDF_AVAILABLE:
            return self._load_pdf_document(path)
        elif file_extension in ['.docx', '.doc'] and DOCX_AVAILABLE:
            return self._load_docx_document(path)
        elif file_extension in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
            return self._load_excel_document(path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _load_text_document(self, path: Path) -> Dict[str, Any]:
        """Load plain text or markdown document."""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            'content': content,
            'format': 'text',
            'metadata': {
                'filename': path.name,
                'size': len(content),
                'format': path.suffix
            }
        }
    
    def _load_pdf_document(self, path: Path) -> Dict[str, Any]:
        """Load PDF document with enhanced text extraction."""
        content = ""
        metadata = {
            'filename': path.name,
            'format': '.pdf',
            'pages': 0,
            'extraction_method': 'pdfplumber'
        }
        
        try:
            # Try pdfplumber first (better for structured documents)
            with pdfplumber.open(path) as pdf:
                metadata['pages'] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n\n"
                        
                # Extract tables if present
                tables = []
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                
                metadata['tables_found'] = len(tables)
                
        except Exception as e:
            self.logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            # Fallback to PyPDF2
            try:
                with open(path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    metadata['extraction_method'] = 'PyPDF2'
                    
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n\n"
                        
            except Exception as e2:
                raise ValueError(f"Failed to parse PDF: {e2}")
        
        metadata['size'] = len(content)
        
        return {
            'content': content,
            'format': 'pdf',
            'metadata': metadata
        }
    
    def _load_docx_document(self, path: Path) -> Dict[str, Any]:
        """Load Word document."""
        try:
            doc = DocxDocument(path)
            content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract tables
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            metadata = {
                'filename': path.name,
                'format': '.docx',
                'size': len(content),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                'content': content,
                'format': 'docx',
                'metadata': metadata,
                'tables': tables_content
            }
            
        except Exception as e:
            raise ValueError(f"Failed to parse Word document: {e}")
    
    def _load_excel_document(self, path: Path) -> Dict[str, Any]:
        """Load Excel document."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(path)
            sheets_content = {}
            content = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(path, sheet_name=sheet_name)
                sheets_content[sheet_name] = df
                
                # Convert to text representation
                content += f"\n\n=== Sheet: {sheet_name} ===\n"
                content += df.to_string(index=False) + "\n"
            
            metadata = {
                'filename': path.name,
                'format': path.suffix,
                'size': len(content),
                'sheets': list(excel_file.sheet_names),
                'total_sheets': len(excel_file.sheet_names)
            }
            
            return {
                'content': content,
                'format': 'excel',
                'metadata': metadata,
                'sheets': sheets_content
            }
            
        except Exception as e:
            raise ValueError(f"Failed to parse Excel document: {e}")
    
    def extract_structured_content(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract structured content from document data.
        
        Args:
            document_data: Document data from load_document()
            
        Returns:
            Structured content with sections, requirements, etc.
        """
        content = document_data['content']
        doc_format = document_data['format']
        
        # Extract sections
        sections = self.extract_sections(content)
        
        # Extract requirements from each section
        all_requirements = []
        for section_code, section_data in sections.items():
            section_requirements = self.extract_requirements(section_data['content'])
            for req in section_requirements:
                all_requirements.append({
                    'section': section_code,
                    'requirement': req
                })
        
        # Extract thresholds and conditions
        thresholds = self.extract_thresholds(content)
        conditions = self.extract_conditions(content)
        
        # Extract metadata specific to policy documents
        policy_metadata = self.extract_policy_metadata(content)
        
        return {
            'sections': sections,
            'requirements': all_requirements,
            'thresholds': thresholds,
            'conditions': conditions,
            'policy_metadata': policy_metadata,
            'document_format': doc_format,
            'original_metadata': document_data['metadata']
        }
    
    def extract_sections(self, document: str) -> Dict[str, Dict[str, str]]:
        """Extract sections from policy document."""
        sections = {}
        
        # Enhanced pattern for various section formats
        patterns = [
            # Standard format: V4.1 OBJECTIVE
            r'(V\d+\.\d+(?:\.\d+)?)\s+([A-Z\s&]+)\n\n(.*?)(?=\n\nV\d+\.\d+|$)',
            # Alternative format: 4.1 Objective
            r'(\d+\.\d+(?:\.\d+)?)\s+([A-Za-z\s&]+)\n\n(.*?)(?=\n\n\d+\.\d+|$)',
            # Header format: ## Section Name
            r'(##\s+)([A-Za-z\s&]+)\n\n(.*?)(?=\n\n##|$)'
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, document, re.DOTALL)
            for match in matches:
                section_code = match.group(1).strip()
                section_title = match.group(2).strip()
                section_content = match.group(3).strip()
                
                sections[section_code] = {
                    'title': section_title,
                    'content': section_content
                }
        
        return sections
    
    def extract_requirements(self, section_content: str) -> List[str]:
        """Extract requirements from section content."""
        requirements = []
        
        # Multiple patterns for different requirement formats
        patterns = [
            # (a), (b), (c) format
            r'\(([a-z]+)\)\s+(.*?)(?=\n\([a-z]+\)|$)',
            # (i), (ii), (iii) format
            r'\(([ivx]+)\)\s+(.*?)(?=\n\([ivx]+\)|$)',
            # 1., 2., 3. format
            r'(\d+)\.\s+(.*?)(?=\n\d+\.|$)',
            # • bullet points
            r'[•·]\s+(.*?)(?=\n[•·]|$)',
            # - dash points
            r'-\s+(.*?)(?=\n-|$)'
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, section_content, re.DOTALL)
            for match in matches:
                requirement = match.group(2).strip() if len(match.groups()) > 1 else match.group(1).strip()
                if requirement and len(requirement) > 10:  # Filter out very short matches
                    requirements.append(requirement)
        
        return requirements
    
    def extract_thresholds(self, document: str) -> Dict[str, Any]:
        """Extract numerical thresholds from document."""
        thresholds = {}
        
        # Enhanced currency extraction
        currency_patterns = [
            r'NZD\s*\$\s*([\d,]+)',
            r'\$\s*([\d,]+)',
            r'([\d,]+)\s*dollars?'
        ]
        
        amounts = []
        for pattern in currency_patterns:
            matches = re.findall(pattern, document, re.IGNORECASE)
            amounts.extend([int(a.replace(',', '')) for a in matches])
        
        # Time periods
        time_pattern = r'(\d+)\s+(months?|years?|days?|weeks?)'
        periods = re.findall(time_pattern, document, re.IGNORECASE)
        
        # Age limits
        age_patterns = [
            r'(?:under|over|age|aged)\s+(\d+)',
            r'(\d+)\s+years?\s+old',
            r'minimum\s+age\s+(\d+)'
        ]
        
        ages = []
        for pattern in age_patterns:
            matches = re.findall(pattern, document, re.IGNORECASE)
            ages.extend([int(a) for a in matches])
        
        thresholds['currency_amounts'] = sorted(set(amounts))
        thresholds['time_periods'] = periods
        thresholds['age_limits'] = sorted(set(ages))
        
        return thresholds
    
    def extract_conditions(self, document: str) -> List[Dict[str, str]]:
        """Extract conditional statements."""
        conditions = []
        
        # Enhanced condition patterns
        condition_patterns = [
            (r'(.*?must\s+.*?)(?:\.|;|\n)', 'mandatory'),
            (r'(.*?shall\s+.*?)(?:\.|;|\n)', 'mandatory'),
            (r'(.*?required\s+to\s+.*?)(?:\.|;|\n)', 'mandatory'),
            (r'(.*?may\s+.*?)(?:\.|;|\n)', 'optional'),
            (r'(.*?can\s+.*?)(?:\.|;|\n)', 'optional'),
            (r'(.*?should\s+.*?)(?:\.|;|\n)', 'recommended'),
            (r'(.*?if\s+.*?)(?:\.|;|\n)', 'conditional')
        ]
        
        for pattern, condition_type in condition_patterns:
            matches = re.finditer(pattern, document, re.IGNORECASE)
            for match in matches:
                statement = match.group(1).strip()
                if len(statement) > 15:  # Filter short matches
                    conditions.append({
                        'type': condition_type,
                        'statement': statement
                    })
        
        return conditions
    
    def extract_policy_metadata(self, document: str) -> Dict[str, Any]:
        """Extract policy-specific metadata."""
        metadata = {}
        
        # Extract visa codes
        visa_codes = re.findall(r'\b[A-Z]\d+\b', document)
        metadata['visa_codes'] = list(set(visa_codes))
        
        # Extract dates
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{1,2}-\d{1,2}-\d{4}',
            r'\d{4}-\d{1,2}-\d{1,2}'
        ]
        
        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, document))
        metadata['dates_mentioned'] = dates
        
        # Extract policy references
        policy_refs = re.findall(r'V\d+\.\d+(?:\.\d+)?(?:\([a-z]+\))?', document)
        metadata['policy_references'] = list(set(policy_refs))
        
        return metadata


# Utility function for easy import
def create_document_parser() -> EnhancedDocumentParser:
    """Create and return an enhanced document parser instance."""
    return EnhancedDocumentParser()


# Check what formats are available
def get_available_formats() -> Dict[str, bool]:
    """Get information about available document formats."""
    return {
        'pdf': PDF_AVAILABLE,
        'docx': DOCX_AVAILABLE,
        'excel': EXCEL_AVAILABLE,
        'text': True,
        'markdown': True
    }
